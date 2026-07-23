# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("情况说明")
run.font.size = Pt(18)
run.font.bold = True
run.font.name = "宋体"

doc.add_paragraph()

today = date.today()
year = today.year
month = today.month
day = today.day

text = (
    f"本人因下班考勤忘记打卡，"
    f"于{year}年{month}月{day}日下午下班未打卡，"
    f"实际已正常出勤，特申请补录考勤。"
)

body = doc.add_paragraph()
body.paragraph_format.first_line_indent = Pt(24)
body.paragraph_format.line_spacing = 1.5
run = body.add_run(text)
run.font.size = Pt(14)
run.font.name = "宋体"

doc.add_paragraph()
doc.add_paragraph()

sign = doc.add_paragraph()
sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sign.add_run("申请人：刘梓彤")
run.font.size = Pt(14)
run.font.name = "宋体"

sign2 = doc.add_paragraph()
sign2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = sign2.add_run(f"日期：{year}年{month}月{day}日")
run.font.size = Pt(14)
run.font.name = "宋体"

filename = "情况说明_刘梓彤.docx"
doc.save(filename)
print("Done: " + filename)
